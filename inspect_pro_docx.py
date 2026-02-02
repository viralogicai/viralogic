import docx
import os
import sys

filename = "Kịch bản EMAIL  ViraLogic PRO.docx"

if not os.path.exists(filename):
    print(f"File not found: {filename}")
    sys.exit(1)

doc = docx.Document(filename)

print(f"Document has {len(doc.paragraphs)} paragraphs.")

full_text = []
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        # print(f"Line {i}: {para.text}") # Commented out to reduce noise, will dump to file
        full_text.append(para.text)

with open("docx_pro_dump.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(full_text))
