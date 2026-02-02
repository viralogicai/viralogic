import docx
import os
import sys

filename = "📩 EMAIL STARTER - 9 Email.docx"
# Handle potential encoding issues with the filename if passed as arg, but here hardcoded
# Check if file exists
if not os.path.exists(filename):
    print(f"File not found: {filename}")
    # Try looking in current dir listing just in case
    print("Files in current dir:")
    for f in os.listdir('.'):
        print(f)
    sys.exit(1)

doc = docx.Document(filename)

print(f"Document has {len(doc.paragraphs)} paragraphs.")

full_text = []
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"Line {i}: {para.text}")
        full_text.append(para.text)

with open("docx_content_dump.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(full_text))
